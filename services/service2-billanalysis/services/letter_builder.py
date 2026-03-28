"""
Letter builder — generates dispute letter in Word (.docx) and PDF formats.
FR-21, FR-22

FR-22 required elements — all must be present or result = system defect:
  - patient name + insurance member ID
  - provider name + identifier
  - session reference ID
  - generation date
  - numbered list of every error (CPT, billed amount, overcharge, regulatory citation)
  - total estimated overcharge
  - formal dispute request paragraph
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors


def build_docx(analysis_data: dict, letter_content: str | None, output_path: str):
    """
    Generate a professionally formatted dispute letter as a Word document.
    Uses python-docx. The Word version is fully editable (US-005 AC4).
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Patient header ────────────────────────────────────────────────────────
    patient_name   = analysis_data.get("patient_name") or "[Patient Name]"
    provider_name  = analysis_data.get("provider_name") or "[Provider Name]"
    date_of_service = analysis_data.get("date_of_service") or "[Date of Service]"
    session_id     = analysis_data.get("session_id", "")
    today          = date.today().strftime("%B %d, %Y")
    total_savings  = analysis_data.get("total_estimated_savings", 0.0)

    # Patient address block
    p = doc.add_paragraph()
    p.add_run(patient_name).bold = True
    doc.add_paragraph("[Address Line 1]")
    doc.add_paragraph("[City, State ZIP]")
    doc.add_paragraph(today)
    doc.add_paragraph()

    # Provider address block
    p = doc.add_paragraph()
    p.add_run(provider_name).bold = True
    doc.add_paragraph("Billing Department")
    doc.add_paragraph("[Provider Address]")
    doc.add_paragraph()

    # Re: line
    p = doc.add_paragraph()
    p.add_run(f"Re: Bill Reference #{session_id[:8].upper()}, "
              f"Date of Service: {date_of_service}").bold = True
    doc.add_paragraph()

    # Salutation
    doc.add_paragraph("Dear Billing Department,")
    doc.add_paragraph()

    # Opening paragraph
    opening = (
        f"I am writing to formally dispute charges on the above-referenced bill "
        f"totalling ${total_savings:,.2f}. After careful review of the itemised "
        f"charges, I have identified the following billing errors:"
    )
    doc.add_paragraph(opening)
    doc.add_paragraph()

    # ── Error list ────────────────────────────────────────────────────────────
    errors = analysis_data.get("errors", [])
    for idx, error in enumerate(errors, start=1):
        cpt_codes = ", ".join(
            str(li) for li in error.get("line_items_affected", [])
        )
        impact    = error.get("estimated_dollar_impact", 0.0)
        error_type = error.get("error_type", "Billing Error")
        description = error.get("description", "")

        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{error_type} — ${impact:,.2f}").bold = True

        doc.add_paragraph(f"    Line item(s) affected: {cpt_codes}")
        doc.add_paragraph(f"    {description}")

        # Citations
        citations = error.get("citations") or []
        if citations:
            for citation in citations:
                source  = citation.get("source", "")
                section = citation.get("section", "")
                doc.add_paragraph(f"    Source: {source}, {section}")

        # RAG explanation (if available)
        explanation = error.get("explanation")
        if explanation:
            doc.add_paragraph(f"    {explanation}")

        doc.add_paragraph()

    # ── Total ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run(f"Total Adjustment Requested: ${total_savings:,.2f}").bold = True
    doc.add_paragraph()

    # ── Formal dispute paragraph ──────────────────────────────────────────────
    if letter_content:
        doc.add_paragraph(letter_content)
    else:
        _add_default_dispute_paragraph(doc, provider_name, total_savings)

    doc.add_paragraph()

    # ── Closing ───────────────────────────────────────────────────────────────
    doc.add_paragraph("I request a written response within 30 days confirming the "
                      "adjustments to be made. Please contact me at the address above "
                      "if you require any additional information.")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(patient_name)
    doc.add_paragraph("[Phone Number]")
    doc.add_paragraph("[Email Address]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"MediCheck Analysis Reference: {session_id}").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run = p.runs[0]
    run.font.size = Pt(8)

    doc.save(output_path)


def build_pdf(analysis_data: dict, letter_content: str | None, output_path: str):
    """
    Generate the same dispute letter as a PDF using ReportLab.
    FR-21: both formats produced and available within the same request.
    """
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=1.25 * inch,
        leftMargin=1.25 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
    )

    styles    = getSampleStyleSheet()
    normal    = styles["Normal"]
    bold_style = ParagraphStyle("Bold", parent=normal, fontName="Helvetica-Bold")
    small      = ParagraphStyle("Small", parent=normal, fontSize=8, textColor=colors.grey)

    patient_name    = analysis_data.get("patient_name") or "[Patient Name]"
    provider_name   = analysis_data.get("provider_name") or "[Provider Name]"
    date_of_service = analysis_data.get("date_of_service") or "[Date of Service]"
    session_id      = analysis_data.get("session_id", "")
    today           = date.today().strftime("%B %d, %Y")
    total_savings   = analysis_data.get("total_estimated_savings", 0.0)
    errors          = analysis_data.get("errors", [])

    story = []

    # Header
    story.append(Paragraph(f"<b>{patient_name}</b>", normal))
    story.append(Paragraph("[Address Line 1]", normal))
    story.append(Paragraph("[City, State ZIP]", normal))
    story.append(Paragraph(today, normal))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(f"<b>{provider_name}</b>", normal))
    story.append(Paragraph("Billing Department", normal))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(
        f"<b>Re: Bill Reference #{session_id[:8].upper()}, "
        f"Date of Service: {date_of_service}</b>", normal
    ))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Dear Billing Department,", normal))
    story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph(
        f"I am writing to formally dispute charges on the above-referenced bill "
        f"totalling <b>${total_savings:,.2f}</b>. After careful review of the itemised "
        f"charges, I have identified the following billing errors:",
        normal
    ))
    story.append(Spacer(1, 0.1 * inch))

    # Error list
    for idx, error in enumerate(errors, start=1):
        impact     = error.get("estimated_dollar_impact", 0.0)
        error_type = error.get("error_type", "Billing Error")
        description = error.get("description", "")
        citations  = error.get("citations") or []
        line_items = ", ".join(str(li) for li in error.get("line_items_affected", []))

        story.append(Paragraph(
            f"<b>{idx}. {error_type} — ${impact:,.2f}</b>", normal
        ))
        story.append(Paragraph(f"Line item(s): {line_items}", normal))
        story.append(Paragraph(description, normal))

        for citation in citations:
            story.append(Paragraph(
                f"Source: {citation.get('source', '')}, {citation.get('section', '')}",
                normal
            ))

        explanation = error.get("explanation")
        if explanation:
            story.append(Paragraph(explanation, normal))

        story.append(Spacer(1, 0.08 * inch))

    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(f"<b>Total Adjustment Requested: ${total_savings:,.2f}</b>", normal))
    story.append(Spacer(1, 0.1 * inch))

    if letter_content:
        story.append(Paragraph(letter_content, normal))
    else:
        story.append(Paragraph(_default_dispute_text(provider_name, total_savings), normal))

    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "I request a written response within 30 days confirming the adjustments "
        "to be made. Please contact me at the address above if you require any "
        "additional information.",
        normal
    ))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Sincerely,", normal))
    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph(f"<b>{patient_name}</b>", normal))
    story.append(Paragraph("[Phone Number]", normal))
    story.append(Paragraph("[Email Address]", normal))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(f"MediCheck Analysis Reference: {session_id}", small))

    doc.build(story)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_default_dispute_paragraph(doc, provider_name: str, total_savings: float):
    text = _default_dispute_text(provider_name, total_savings)
    doc.add_paragraph(text)


def _default_dispute_text(provider_name: str, total_savings: float) -> str:
    return (
        f"I respectfully request that {provider_name} review the above itemised "
        f"errors and issue a corrected bill reflecting a total adjustment of "
        f"${total_savings:,.2f}. These errors were identified using AI-assisted "
        f"billing analysis cross-referenced against the CMS Physician Fee Schedule "
        f"and applicable federal regulations including the No Surprises Act "
        f"(Pub. L. 116-260). I am prepared to escalate this dispute to my state "
        f"insurance commissioner if a satisfactory resolution is not reached within "
        f"30 days of this letter."
    )
